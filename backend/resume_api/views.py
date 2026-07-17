from django.shortcuts import render
from rest_framework import status, viewsets, mixins
from rest_framework.decorators import api_view, parser_classes, permission_classes, throttle_classes, action
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.throttling import ScopedRateThrottle
from django.conf import settings
from django.contrib.auth.models import User
from django.utils import timezone
from django.views.decorators.cache import cache_page
import os
import json
import tempfile
import uuid

from .pagination import ActivityPagination

# Uploaded files are read fully into memory for text extraction, so cap
# their size well below what would strain the 512MB Render free-tier worker.
MAX_UPLOAD_SIZE_BYTES = 10 * 1024 * 1024  # 10MB
ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}
ALLOWED_AUDIO_EXTENSIONS = {'wav', 'webm', 'mp3', 'ogg', 'm4a'}


def _validate_upload(file_obj, allowed_extensions, label):
    """Returns an error string if the upload fails size/type checks, else None."""
    if file_obj.size > MAX_UPLOAD_SIZE_BYTES:
        return f"{label} is too large ({file_obj.size // (1024 * 1024)}MB). Maximum size is 10MB."
    ext = file_obj.name.rsplit('.', 1)[-1].lower() if '.' in file_obj.name else ''
    if ext not in allowed_extensions:
        return f"{label} has an unsupported file type '.{ext}'. Allowed types: {', '.join(sorted(allowed_extensions))}."
    return None

from .models import Resume, JobDescription, ResumeAnalysis, ChatMessage, MockInterview, UserActivity, BookmarkedAnswer, Notification
from .serializers import (
    ResumeSerializer,
    JobDescriptionSerializer,
    ResumeAnalysisSerializer,
    ResumeAnalysisResultSerializer,
    MockInterviewSerializer,
    InterviewAnalysisResultSerializer,
    InterviewFeedbackSerializer,
    ChatMessageSerializer,
    UserActivitySerializer,
    BookmarkedAnswerSerializer,
    NotificationSerializer,
)
from .resume_analyzer import ResumeAnalyzer
from .interview_analyzer import InterviewAnalyzer
from . import azure_language_client
from . import azure_speech_client
from . import azure_language_client
from .groq_client import InterviewChatbot, get_groq_client
import json

# Initialize the resume analyzer and interview chatbot
resume_analyzer = ResumeAnalyzer()
interview_chatbot = InterviewChatbot()
interview_analyzer = InterviewAnalyzer()

@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def analyze_resume(request):
    """
    Analyze a resume against a job description and provide tailoring suggestions.
    """
    resume_file = request.FILES.get('resume_file')
    job_desc_file = request.FILES.get('job_desc_file')
    industry = request.data.get('industry') or None

    if not resume_file or not job_desc_file:
        return Response({
            'error': 'Both resume and job description files are required.'
        }, status=status.HTTP_400_BAD_REQUEST)

    for upload, label in ((resume_file, 'Resume'), (job_desc_file, 'Job description')):
        validation_error = _validate_upload(upload, ALLOWED_DOCUMENT_EXTENSIONS, label)
        if validation_error:
            return Response({'error': validation_error}, status=status.HTTP_400_BAD_REQUEST)

    analyzer = ResumeAnalyzer()
    
    # Extract text from files
    resume_text = analyzer.extract_text_from_file(
        resume_file.read(), 
        resume_file.name.split('.')[-1]
    )
    
    job_desc_text = analyzer.extract_text_from_file(
        job_desc_file.read(), 
        job_desc_file.name.split('.')[-1]
    )
    
    # Analyze the resume against the job description
    analysis_result = analyzer.analyze_resume_and_job_description(
        resume_text,
        job_desc_text,
        industry=industry
    )
    
    # Log the complete results for debugging
    print("Complete analysis result:", json.dumps(analysis_result, default=str, indent=2))
    
    # Specifically check the sentiment analysis part
    sentiment_data = analysis_result.get('sentimentAnalysis', {})
    print("Sentiment Analysis data:", json.dumps(sentiment_data, default=str, indent=2))

    # Include the extracted resume text so the frontend can render an inline
    # diff of the suggested edits without re-uploading the file.
    analysis_result['resumeText'] = resume_text

    # Persist the analysis for signed-in users so it shows up in their
    # history (via /api/resume/analyses/) instead of disappearing once the
    # page is closed — mirrors how mock interviews are already saved.
    if request.user.is_authenticated:
        try:
            resume = Resume.objects.create(
                user=request.user,
                title=resume_file.name,
                file_name=resume_file.name,
                content=resume_text,
                file_type=resume_file.name.split('.')[-1],
            )
            job_description = JobDescription.objects.create(
                user=request.user,
                title=job_desc_file.name,
                file_name=job_desc_file.name,
                content=job_desc_text,
                file_type=job_desc_file.name.split('.')[-1],
            )
            saved_analysis = ResumeAnalysis.objects.create(
                user=request.user,
                resume=resume,
                job_description=job_description,
                keywords_to_add=analysis_result.get('keywordsToAdd', []),
                keywords_to_remove=analysis_result.get('keywordsToRemove', []),
                format_suggestions=analysis_result.get('formatSuggestions', []),
                content_suggestions=analysis_result.get('contentSuggestions', []),
                match_score=analysis_result.get('matchScore', 0),
            )
            # Lets the frontend immediately offer a "Download tailored resume"
            # button without a round trip through the analysis history list.
            analysis_result['analysisId'] = saved_analysis.id
        except Exception as e:
            # Saving history is a convenience feature — never let a DB hiccup
            # break the analysis response the user is actively waiting on.
            print(f"Failed to persist resume analysis history: {e}")

    return Response(analysis_result, status=status.HTTP_200_OK)

@api_view(['POST'])
def test_sentiment_analysis(request):
    """
    Debug endpoint to test sentiment analysis functionality.
    """
    text = request.data.get('text', 'This is a test text with a neutral sentiment.')
    result = azure_language_client.analyze_sentiment(text)
    
    # Log the raw result
    print("Raw sentiment analysis result:", result)
    
    # Return the full result
    return Response(result, status=status.HTTP_200_OK)

@api_view(['POST'])
@parser_classes([MultiPartParser, FormParser])
def analyze_interview(request):
    """
    Analyze a mock interview recording and provide feedback.
    """
    audio_file = request.FILES.get('audio_file')
    job_description_id = request.data.get('job_description_id')
    interview_title = request.data.get('title', 'Mock Interview')
    
    if not audio_file:
        return Response({
            'error': 'Audio file is required.'
        }, status=status.HTTP_400_BAD_REQUEST)

    validation_error = _validate_upload(audio_file, ALLOWED_AUDIO_EXTENSIONS, 'Audio recording')
    if validation_error:
        return Response({'error': validation_error}, status=status.HTTP_400_BAD_REQUEST)

    # Get job description text if provided
    job_description_text = ""
    job_description = None
    
    if job_description_id:
        try:
            job_description = JobDescription.objects.get(id=job_description_id)
            job_description_text = job_description.content
        except JobDescription.DoesNotExist:
            return Response({
                'error': 'Job description not found.'
            }, status=status.HTTP_404_NOT_FOUND)
    
    # Read audio data from file
    audio_data = audio_file.read()
    
    # Save audio file for later reference
    audio_file_path = None
    if audio_data:
        # Create a unique filename
        filename = f"interview_{uuid.uuid4().hex}.wav"
        audio_dir = os.path.join(settings.MEDIA_ROOT, 'interviews')
        
        # Create directory if it doesn't exist
        os.makedirs(audio_dir, exist_ok=True)
        
        # Save audio file
        audio_file_path = os.path.join(audio_dir, filename)
        with open(audio_file_path, 'wb') as f:
            f.write(audio_data)
    
    # Analyze the interview
    analysis_result = interview_analyzer.analyze_interview(audio_data, job_description_text)
    
    if "error" in analysis_result:
        return Response({
            'error': analysis_result["error"]
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Create or update the MockInterview model instance
    mock_interview = MockInterview(
        user=request.user,
        job_description=job_description,
        title=interview_title,
        transcript=analysis_result.get("transcript", ""),
        audio_file_path=audio_file_path,
        duration=analysis_result.get("audio_analysis", {}).get("duration", 0),
        audio_analysis=analysis_result.get("audio_analysis", {}),
        content_analysis=analysis_result.get("content_analysis", {}),
        feedback=analysis_result.get("feedback", {}),
        overall_score=analysis_result.get("feedback", {}).get("overall_score", {}).get("score", 0)
    )
    mock_interview.save()
    
    # Return the analysis result
    serializer = InterviewAnalysisResultSerializer(analysis_result)
    return Response(serializer.data, status=status.HTTP_200_OK)

@api_view(['GET'])
def get_interview_feedback(request, interview_id):
    """
    Get detailed feedback for a specific mock interview.
    """
    try:
        interview = MockInterview.objects.get(id=interview_id, user=request.user)
        feedback = interview.feedback
        
        serializer = InterviewFeedbackSerializer(feedback)
        return Response(serializer.data, status=status.HTTP_200_OK)
    except MockInterview.DoesNotExist:
        return Response({
            'error': 'Mock interview not found.'
        }, status=status.HTTP_404_NOT_FOUND)
    
@api_view(['POST'])
@permission_classes([AllowAny])
@throttle_classes([ScopedRateThrottle])
def interview_chat(request):
    """
    Process a user message to the interview preparation chatbot and get a response.
    """
    try:
        user_message = request.data.get('message', '')
        session_id = request.data.get('session_id', '')
        mode = request.data.get('mode', 'advice')  # 'advice' or 'reverse' (chatbot interviews you)
        job_posting = request.data.get('job_posting', '') or None
        want_follow_ups = bool(request.data.get('want_follow_ups', False))
        user_id = request.user.id if request.user.is_authenticated else None

        if not user_message:
            return Response({
                'error': 'Message is required.'
            }, status=status.HTTP_400_BAD_REQUEST)

        # Generate a new session ID if none provided
        if not session_id:
            session_id = interview_chatbot.generate_session_id()

        # Get previous messages in this session if user is authenticated
        previous_messages = []
        if user_id:
            previous_messages = ChatMessage.objects.filter(
                user_id=user_id,
                session_id=session_id
            ).order_by('timestamp')

        # Save the user message if user is authenticated
        if user_id:
            user_chat_message = ChatMessage.objects.create(
                user_id=user_id,
                message=user_message,
                is_user=True,
                session_id=session_id
            )

        # Get response from the chatbot with better error handling
        print(f"Sending message to Groq: '{user_message}'")
        bot_response = interview_chatbot.get_response(
            user_message, previous_messages, mode=mode, job_posting=job_posting
        )
        print(f"Received response from Groq: '{bot_response}'")

        # Save the bot response if user is authenticated
        if user_id:
            bot_chat_message = ChatMessage.objects.create(
                user_id=user_id,
                message=bot_response,
                is_user=False,
                session_id=session_id
            )

        follow_up_questions = []
        if want_follow_ups:
            follow_up_questions = interview_chatbot.get_follow_up_questions(user_message, bot_response)

        # Lightweight usage-quota indicator mirroring the 'interview_chat'
        # throttle scope's own 30/minute limit (see settings.DEFAULT_THROTTLE_RATES),
        # so the frontend can show "X messages left this minute" without
        # reaching into DRF's internal throttle cache format.
        from django.core.cache import cache
        quota_ident = f"user_{user_id}" if user_id else f"ip_{request.META.get('REMOTE_ADDR', 'unknown')}"
        quota_key = f"interview_chat_quota:{quota_ident}"
        quota_count = cache.get(quota_key, 0) + 1
        cache.set(quota_key, quota_count, timeout=60)
        quota_limit = 30
        quota_remaining = max(0, quota_limit - quota_count)

        return Response({
            'message': bot_response,
            'session_id': session_id,
            'follow_up_questions': follow_up_questions,
            'quota': {'limit': quota_limit, 'remaining': quota_remaining, 'resets_in_seconds': 60},
        }, status=status.HTTP_200_OK)

    except Exception as e:
        print(f"Error in interview_chat view: {str(e)}")
        return Response({
            'error': 'An unexpected error occurred.',
            'details': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# @api_view wraps this in a generated APIView subclass exposed as `.cls`;
# throttle_scope must be set there (not on the function) for ScopedRateThrottle
# to find it — this is the public, unauthenticated endpoint most exposed to abuse.
interview_chat.cls.throttle_scope = 'interview_chat'

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_chat_history(request):
    """
    Get the chat history for a specific session.
    """
    session_id = request.query_params.get('session_id', '')
    
    if not session_id:
        return Response({
            'error': 'Session ID is required.'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    # Get messages for this session
    messages = ChatMessage.objects.filter(
        user=request.user,
        session_id=session_id
    ).order_by('timestamp')
    
    serializer = ChatMessageSerializer(messages, many=True)
    
    return Response(serializer.data, status=status.HTTP_200_OK)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_chat_sessions(request):
    """
    Get a list of all chat sessions for the authenticated user.
    """
    # Get all unique session IDs for this user. ChatMessage's default
    # ordering (Meta.ordering = ['timestamp']) gets pulled into the SELECT
    # for DISTINCT to be well-defined, which silently turns this into
    # "distinct (session_id, timestamp)" and returns the same session once
    # per message. order_by() clears that default ordering first.
    sessions = ChatMessage.objects.filter(
        user=request.user
    ).order_by().values('session_id').distinct()
    
    session_data = []
    for session in sessions:
        session_id = session['session_id']
        # Get the first message in this session
        first_message = ChatMessage.objects.filter(
            user=request.user,
            session_id=session_id,
            is_user=True
        ).order_by('timestamp').first()
        
        # Get the timestamp of the last message
        last_message = ChatMessage.objects.filter(
            user=request.user,
            session_id=session_id
        ).order_by('-timestamp').first()
        
        if first_message and last_message:
            session_data.append({
                'session_id': session_id,
                'title': first_message.message[:50] + ('...' if len(first_message.message) > 50 else ''),
                'last_updated': last_message.timestamp
            })
    
    # Sort sessions by last updated timestamp, newest first
    session_data.sort(key=lambda x: x['last_updated'], reverse=True)
    
    return Response(session_data, status=status.HTTP_200_OK)

@cache_page(60 * 60)
@api_view(['GET'])
@permission_classes([AllowAny])
def get_faq_topics(request):
    """
    Get a list of predefined FAQ topics for the interview preparation chatbot.
    """
    faq_topics = list(interview_chatbot.predefined_faqs.keys())
    
    return Response({
        'topics': faq_topics
    }, status=status.HTTP_200_OK)

class ResumeViewSet(viewsets.ModelViewSet):
    """ViewSet for viewing and editing Resume instances"""
    serializer_class = ResumeSerializer
    
    def get_queryset(self):
        """Return resumes for the current authenticated user only"""
        user = self.request.user
        if user.is_authenticated:
            return Resume.objects.filter(user=user).order_by('-created_at')
        return Resume.objects.none()
    
    def perform_create(self, serializer):
        """Set the user when creating a new resume"""
        serializer.save(user=self.request.user)

class JobDescriptionViewSet(viewsets.ModelViewSet):
    """ViewSet for viewing and editing JobDescription instances"""
    serializer_class = JobDescriptionSerializer
    
    def get_queryset(self):
        """Return job descriptions for the current authenticated user only"""
        user = self.request.user
        if user.is_authenticated:
            return JobDescription.objects.filter(user=user).order_by('-created_at')
        return JobDescription.objects.none()
    
    def perform_create(self, serializer):
        """Set the user when creating a new job description"""
        serializer.save(user=self.request.user)

class ResumeAnalysisViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing ResumeAnalysis instances"""
    serializer_class = ResumeAnalysisSerializer
    
    def get_queryset(self):
        """Return analyses for the current authenticated user only"""
        user = self.request.user
        if user.is_authenticated:
            return ResumeAnalysis.objects.filter(user=user).order_by('-created_at')
        return ResumeAnalysis.objects.none()

class MockInterviewViewSet(viewsets.ModelViewSet):
    """ViewSet for viewing and managing MockInterview instances"""
    serializer_class = MockInterviewSerializer
    
    def get_queryset(self):
        """Return mock interviews for the current authenticated user only"""
        user = self.request.user
        if user.is_authenticated:
            return MockInterview.objects.filter(user=user).order_by('-created_at')
        return MockInterview.objects.none()
    
    def perform_create(self, serializer):
        """Set the user when creating a new mock interview"""
        serializer.save(user=self.request.user)

class UserActivityViewSet(viewsets.ModelViewSet):
    """
    Per-account dashboard activity log (stats, score chart, achievements,
    recent-activity feed). Every request is scoped to request.user, so this
    data lives with the account, not the browser/device.
    """
    serializer_class = UserActivitySerializer
    pagination_class = ActivityPagination

    def get_queryset(self):
        user = self.request.user
        if user.is_authenticated:
            return UserActivity.objects.filter(user=user).order_by('-created_at')
        return UserActivity.objects.none()

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

# Each entry is a (label, system prompt) pair used to steer the Groq model
# toward one specific career-coaching task. Keyed by the id the frontend
# sends so new tools can be added here without touching the view logic.
AI_TOOLS = {
    'salary_negotiation': (
        'Salary Negotiation Coach',
        "You are a salary negotiation coach. Given the user's job offer details and goals, "
        "provide a short negotiation script they could say out loud, plus 3-4 concise tips. "
        "Be practical and specific to the numbers/context they gave you."
    ),
    'linkedin_headline': (
        'LinkedIn Headline & Summary Optimizer',
        "You are a LinkedIn profile optimization expert. Given the user's role, experience, and "
        "goals, write 3 alternative LinkedIn headlines (under 220 characters each) and a 3-sentence "
        "'About' summary. Keep it professional and specific, not generic buzzwords."
    ),
    'job_red_flags': (
        'Job Posting Red-Flag Detector',
        "You are a job posting analyst. Given a job description, list any red flags (vague pay, "
        "unrealistic requirements for the level, excessive scope, culture warning signs) as bullet "
        "points, then give an overall risk rating of Low, Medium, or High with a one-line reason."
    ),
    'networking_message': (
        'Networking Message Generator',
        "You are a professional networking coach. Given who the user wants to contact and why, "
        "write a concise, warm, non-pushy outreach message (suitable for LinkedIn or email) under "
        "150 words."
    ),
    'skill_gap': (
        'Skill Gap Analyzer',
        "You are a career coach. Given the user's current skills and their target role, list the "
        "key skill gaps as bullet points, and for each one suggest a specific type of course, "
        "certification, or project that would close it."
    ),
    'company_research': (
        'Company Research Briefing',
        "You are a company research assistant preparing a candidate for an interview. Given the "
        "company name and any details provided, produce a short briefing: likely culture/values, "
        "what they probably care about right now, and 3 smart questions to ask the interviewer."
    ),
    'elevator_pitch': (
        'Elevator Pitch Generator',
        "You are an elevator pitch coach. Given the user's background and the role they're "
        "targeting, write a natural-sounding 30-second spoken elevator pitch (roughly 75-100 words)."
    ),
    'thank_you_note': (
        'Post-Interview Thank-You Note',
        "You are an interview follow-up coach. Given details about the interview (role, interviewer, "
        "topics discussed), write a concise, warm, professional thank-you email, under 150 words."
    ),
    'interview_predictor': (
        'Interview Question Predictor',
        "You are an interview question predictor. Given a job description, predict the 6-8 most "
        "likely interview questions for this role, grouped under headings like Behavioral, "
        "Technical, and Situational."
    ),
    'bullet_rewriter': (
        'Resume Bullet-Point Rewriter',
        "You are a resume writing coach. Given the user's raw job duties or bullet points, rewrite "
        "each one into a concise, achievement-oriented bullet point that leads with a strong action "
        "verb and includes a quantified result wherever plausible. Return one rewritten bullet per "
        "line, in the same order as the input."
    ),
    'cover_letter': (
        'Cover Letter Generator',
        "You are a cover letter writer. Given the user's background/resume summary and the job "
        "description or role they're applying to, write a concise, tailored cover letter (3-4 short "
        "paragraphs) that connects their experience to the role without simply repeating the resume."
    ),
    'ats_checker': (
        'ATS Compatibility Checker',
        "You are an ATS (Applicant Tracking System) compatibility expert. Given resume text, check for "
        "formatting and content issues that commonly break ATS parsing: tables/columns, headers/footers, "
        "images or icons instead of text, non-standard section headings, unusual fonts/symbols, and "
        "missing standard sections (Experience, Education, Skills). List each issue found as a bullet, "
        "then give an overall ATS-friendliness rating of Good, Fair, or Poor with a one-line reason."
    ),
}


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@throttle_classes([ScopedRateThrottle])
def ai_tools_generate(request):
    """
    Generic endpoint powering the AI Tools hub (salary coach, LinkedIn
    optimizer, red-flag detector, etc.) — one Groq call per request, steered
    by a per-tool system prompt, so each new tool is just a dictionary entry
    rather than a new view.
    """
    tool = request.data.get('tool', '')
    user_input = request.data.get('input', '').strip()

    if tool not in AI_TOOLS:
        return Response({'error': 'Unknown tool.'}, status=status.HTTP_400_BAD_REQUEST)
    if not user_input:
        return Response({'error': 'Please provide some input for this tool.'}, status=status.HTTP_400_BAD_REQUEST)

    label, system_prompt = AI_TOOLS[tool]

    try:
        client = get_groq_client()
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_input},
            ],
            temperature=0.7,
            max_tokens=700,
        )
        result = response.choices[0].message.content.strip()
        return Response({'tool': tool, 'label': label, 'result': result}, status=status.HTTP_200_OK)
    except Exception as e:
        return Response(
            {'error': 'The AI service is temporarily unavailable. Please try again in a moment.'},
            status=status.HTTP_503_SERVICE_UNAVAILABLE,
        )

ai_tools_generate.cls.throttle_scope = 'ai_tools'

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def save_interview_attempt(request):
    """
    Save a mock interview attempt analyzed client-side (see clientAnalysisService.ts)
    so it shows up in the user's history. MockInterviewSerializer marks
    overall_score read-only (it's normally set by the server-side Azure
    analysis path), so this bypasses that serializer entirely — the score
    here is the user's own client-computed result for their own recording,
    not something that needs protecting from them.
    """
    title = request.data.get('title', 'Mock Interview')
    transcript = request.data.get('transcript', '')
    duration = request.data.get('duration', 0)
    overall_score = request.data.get('overall_score', 0)
    difficulty = request.data.get('difficulty', 'mid')
    question_text = request.data.get('question_text', '')
    session_id = request.data.get('session_id', '')

    try:
        overall_score = int(overall_score)
        duration = float(duration)
    except (TypeError, ValueError):
        return Response({'error': 'overall_score and duration must be numeric.'}, status=status.HTTP_400_BAD_REQUEST)

    if difficulty not in dict(MockInterview.DIFFICULTY_CHOICES):
        difficulty = 'mid'

    mock_interview = MockInterview.objects.create(
        user=request.user,
        title=title,
        transcript=transcript,
        duration=duration,
        overall_score=overall_score,
        difficulty=difficulty,
        question_text=question_text,
        session_id=session_id,
        audio_analysis=request.data.get('audio_analysis', {}) or {},
        content_analysis=request.data.get('content_analysis', {}) or {},
        feedback=request.data.get('feedback', {}) or {},
    )

    # Achievement-style notification for a new personal best score.
    previous_best = (
        MockInterview.objects.filter(user=request.user)
        .exclude(id=mock_interview.id)
        .order_by('-overall_score')
        .values_list('overall_score', flat=True)
        .first()
    )
    if previous_best is not None and overall_score > previous_best:
        notify_user(
            request.user, 'achievement', 'New personal best!',
            f"You scored {overall_score} on your latest mock interview, beating your previous best of {previous_best}."
        )

    return Response(MockInterviewSerializer(mock_interview).data, status=status.HTTP_201_CREATED)

class BookmarkedAnswerViewSet(viewsets.ModelViewSet):
    """Chatbot answers the user starred for quick reference later."""
    serializer_class = BookmarkedAnswerSerializer

    def get_queryset(self):
        user = self.request.user
        if user.is_authenticated:
            return BookmarkedAnswer.objects.filter(user=user).order_by('-created_at')
        return BookmarkedAnswer.objects.none()

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class NotificationViewSet(
    mixins.ListModelMixin,
    mixins.RetrieveModelMixin,
    mixins.UpdateModelMixin,
    mixins.DestroyModelMixin,
    viewsets.GenericViewSet,
):
    """
    Notification center backing the dashboard bell icon (streaks,
    achievements, digest summaries). Read-mostly from the frontend's
    perspective — creation happens server-side (see notify_user below),
    never via a POST from the client (this viewset has no create route).
    """
    serializer_class = NotificationSerializer

    def get_queryset(self):
        user = self.request.user
        if user.is_authenticated:
            return Notification.objects.filter(user=user)
        return Notification.objects.none()

    @action(detail=False, methods=['post'])
    def mark_all_read(self, request):
        self.get_queryset().filter(is_read=False).update(is_read=True)
        return Response({'detail': 'All notifications marked as read.'}, status=status.HTTP_200_OK)


def notify_user(user, notification_type, title, message):
    """Create an in-app notification, respecting the user's per-type opt-out preferences."""
    pref_field = {
        'achievement': 'notify_achievement_alerts',
        'streak': 'notify_streak_reminders',
        'digest': 'notify_progress_digest',
    }.get(notification_type)
    if pref_field and not getattr(user.profile, pref_field, True):
        return None
    return Notification.objects.create(
        user=user, notification_type=notification_type, title=title, message=message
    )


@cache_page(60 * 60)
@api_view(['GET'])
@permission_classes([AllowAny])
def keyword_libraries(request):
    """Industry-specific keyword libraries available for resume analysis."""
    from .keyword_libraries import INDUSTRY_KEYWORDS
    return Response({
        'industries': [
            {'id': key, 'label': key.title(), 'keywords': keywords}
            for key, keywords in INDUSTRY_KEYWORDS.items()
        ]
    }, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_tailored_resume(request, analysis_id):
    """
    Reconstruct a downloadable .docx from the original resume text plus the
    analysis's suggestions — the original content stays intact and untouched
    (so nothing is silently rewritten wrong), with a clearly labeled
    "Tailoring Suggestions" section appended covering keywords to add,
    content suggestions, and formatting fixes to apply by hand.
    """
    import io
    import docx
    from django.http import FileResponse

    try:
        analysis = ResumeAnalysis.objects.select_related('resume', 'job_description').get(
            id=analysis_id, user=request.user
        )
    except ResumeAnalysis.DoesNotExist:
        return Response({'error': 'Analysis not found.'}, status=status.HTTP_404_NOT_FOUND)

    document = docx.Document()
    document.add_heading(analysis.resume.title or 'Tailored Resume', level=1)
    document.add_paragraph(f"Tailored for: {analysis.job_description.title}")
    document.add_paragraph(f"Match score: {analysis.match_score}%")
    document.add_paragraph('')

    document.add_heading('Original Resume Content', level=2)
    for line in analysis.resume.content.splitlines():
        document.add_paragraph(line if line.strip() else '')

    document.add_page_break()
    document.add_heading('Tailoring Suggestions to Apply', level=2)

    if analysis.keywords_to_add:
        document.add_heading('Keywords to Add', level=3)
        for kw in analysis.keywords_to_add:
            document.add_paragraph(kw, style='List Bullet')

    if analysis.keywords_to_remove:
        document.add_heading('Keywords to Reconsider Removing', level=3)
        for kw in analysis.keywords_to_remove:
            document.add_paragraph(kw, style='List Bullet')

    if analysis.content_suggestions:
        document.add_heading('Content Suggestions', level=3)
        for suggestion in analysis.content_suggestions:
            document.add_paragraph(suggestion, style='List Bullet')

    if analysis.format_suggestions:
        document.add_heading('Formatting Suggestions', level=3)
        for suggestion in analysis.format_suggestions:
            document.add_paragraph(suggestion, style='List Bullet')

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"tailored-{(analysis.resume.title or 'resume').replace(' ', '-')}.docx"
    return FileResponse(
        buffer, as_attachment=True, filename=filename,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_share_link(request, interview_id):
    """Generate (or return the existing) shareable read-only link for a mock interview attempt."""
    try:
        interview = MockInterview.objects.get(id=interview_id, user=request.user)
    except MockInterview.DoesNotExist:
        return Response({'error': 'Mock interview not found.'}, status=status.HTTP_404_NOT_FOUND)

    if not interview.share_token:
        interview.share_token = uuid.uuid4()
        interview.save(update_fields=['share_token'])

    return Response({'share_token': str(interview.share_token)}, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([AllowAny])
def shared_interview_view(request, share_token):
    """Public, read-only view of a mock interview attempt's feedback — for sharing with a mentor."""
    try:
        interview = MockInterview.objects.get(share_token=share_token)
    except (MockInterview.DoesNotExist, ValueError):
        return Response({'error': 'This share link is invalid or has expired.'}, status=status.HTTP_404_NOT_FOUND)

    return Response({
        'title': interview.title,
        'difficulty': interview.difficulty,
        'question_text': interview.question_text,
        'transcript': interview.transcript,
        'overall_score': interview.overall_score,
        'feedback': interview.feedback,
        'created_at': interview.created_at,
    }, status=status.HTTP_200_OK)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@throttle_classes([ScopedRateThrottle])
def adaptive_follow_up(request):
    """Generate one adaptive follow-up question probing deeper into the candidate's last answer."""
    question = request.data.get('question', '')
    transcript = request.data.get('transcript', '')
    if not question or not transcript:
        return Response({'error': 'question and transcript are required.'}, status=status.HTTP_400_BAD_REQUEST)

    follow_up = interview_chatbot.get_adaptive_follow_up(question, transcript)
    return Response({'follow_up_question': follow_up}, status=status.HTTP_200_OK)

adaptive_follow_up.cls.throttle_scope = 'interview_chat'


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def question_score_trend(request):
    """
    Score trend for repeated attempts at the same question — powers the
    "have I improved on this specific question?" chart in interview history.
    Only questions attempted 2+ times are returned (a trend needs >1 point).
    """
    from django.db.models import Count

    interviews = (
        MockInterview.objects.filter(user=request.user)
        .exclude(question_text='')
        .order_by('question_text', 'created_at')
    )

    by_question = {}
    for interview in interviews:
        by_question.setdefault(interview.question_text, []).append({
            'id': interview.id,
            'score': interview.overall_score,
            'created_at': interview.created_at,
        })

    trends = [
        {'question_text': question, 'attempts': attempts}
        for question, attempts in by_question.items()
        if len(attempts) >= 2
    ]
    trends.sort(key=lambda t: len(t['attempts']), reverse=True)

    return Response({'trends': trends}, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def leaderboard(request):
    """
    Opt-in, anonymized leaderboard of best mock-interview scores. Only users
    with profile.leaderboard_opt_in=True are included — everyone else is
    invisible to this endpoint entirely, not just unranked. Usernames are
    masked (first 2 chars + asterisks) since this is meant to be a fun,
    low-stakes comparison, not a public identity list.
    """
    from django.db.models import Max
    from users.models import Profile

    opted_in_user_ids = Profile.objects.filter(leaderboard_opt_in=True).values_list('user_id', flat=True)
    best_scores = (
        MockInterview.objects.filter(user_id__in=opted_in_user_ids)
        .values('user_id', 'user__username')
        .annotate(best_score=Max('overall_score'))
        .order_by('-best_score')[:20]
    )

    def mask(username):
        return (username[:2] + '*' * max(0, len(username) - 2)) if username else 'Anonymous'

    entries = [
        {
            'rank': i + 1,
            'display_name': mask(row['user__username']),
            'best_score': row['best_score'],
            'is_you': row['user_id'] == request.user.id,
        }
        for i, row in enumerate(best_scores)
    ]

    return Response({'entries': entries, 'you_opted_in': request.user.id in set(opted_in_user_ids)}, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def benchmark(request):
    """
    Anonymized comparison of the current user's average mock-interview and
    resume-match scores against the platform-wide average. Unlike the
    leaderboard, this doesn't require opting in — it only ever reveals an
    aggregate average, never another individual's score.
    """
    from django.db.models import Avg

    user_interview_avg = MockInterview.objects.filter(user=request.user).aggregate(avg=Avg('overall_score'))['avg']
    platform_interview_avg = MockInterview.objects.aggregate(avg=Avg('overall_score'))['avg']

    user_resume_avg = ResumeAnalysis.objects.filter(user=request.user).aggregate(avg=Avg('match_score'))['avg']
    platform_resume_avg = ResumeAnalysis.objects.aggregate(avg=Avg('match_score'))['avg']

    def round_or_none(value):
        return round(value, 1) if value is not None else None

    return Response({
        'mock_interview': {
            'your_average': round_or_none(user_interview_avg),
            'platform_average': round_or_none(platform_interview_avg),
        },
        'resume_match': {
            'your_average': round_or_none(user_resume_avg),
            'platform_average': round_or_none(platform_resume_avg),
        },
    }, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def next_best_action(request):
    """
    A single, personalized "what should I do next" suggestion for the
    dashboard, derived from simple rules over the user's own data — no ML,
    just prioritized heuristics so the suggestion is always explainable.
    """
    user = request.user
    resume_count = Resume.objects.filter(user=user).count()
    interview_count = MockInterview.objects.filter(user=user).count()
    latest_analysis = ResumeAnalysis.objects.filter(user=user).order_by('-created_at').first()
    last_activity = UserActivity.objects.filter(user=user).order_by('-created_at').first()

    if resume_count == 0:
        action = {
            'title': 'Upload your first resume',
            'description': 'Get tailored suggestions by analyzing your resume against a job description.',
            'cta_label': 'Analyze a resume',
            'cta_path': '/resume',
        }
    elif latest_analysis and latest_analysis.match_score < 70:
        action = {
            'title': 'Improve your resume match score',
            'description': f"Your latest resume scored {latest_analysis.match_score}% against the job description. Apply the suggested keywords to boost it.",
            'cta_label': 'View suggestions',
            'cta_path': '/resume',
        }
    elif interview_count == 0:
        action = {
            'title': 'Try a mock interview',
            'description': 'Practice answering real interview questions and get instant feedback.',
            'cta_label': 'Start mock interview',
            'cta_path': '/mock-interview',
        }
    elif last_activity is None or (timezone.now() - last_activity.created_at).days >= 3:
        action = {
            'title': "You've been away — jump back in",
            'description': 'Keep your streak going with a quick mock interview or chatbot session.',
            'cta_label': 'Practice now',
            'cta_path': '/mock-interview',
        }
    else:
        action = {
            'title': 'Explore the AI Career Tools',
            'description': 'Try the elevator pitch generator, salary negotiation coach, or skill-gap analyzer.',
            'cta_label': 'Open AI Tools',
            'cta_path': '/ai-tools',
        }

    return Response(action, status=status.HTTP_200_OK)
